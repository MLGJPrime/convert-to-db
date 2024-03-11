from docx import Document


def convert_word_to_txt(word_file_path, txt_file_path):
    doc = Document(word_file_path)
    output = []

    for para in doc.paragraphs:
        for run in para.runs:
            if run.bold:
                output.append(f"*{run.text}*")
            else:
                output.append(run.text)

    with open(txt_file_path, 'w') as f:
        f.write('\n'.join(output))


def convert_txt_to_word(txt_file_path, word_file_path):
    with open(txt_file_path, 'r') as f:
        lines = f.read().split('\n')

    doc = Document()

    for line in lines:
        para = doc.add_paragraph()
        if line.startswith('*') and line.endswith('*'):
            run = para.add_run(line[1:-1])  # remove the enclosing *
            run.bold = True
        else:
            para.add_run(line)

    doc.save(word_file_path)


def main():
    # Usage
    convert_word_to_txt('french.docx', 'french.txt')
    convert_txt_to_word('german.txt', 'german.docx')


if __name__ == "__main__":
    main()
